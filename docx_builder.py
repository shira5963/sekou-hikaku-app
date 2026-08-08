"""
施工前・施工後 比較報告書 Word(.docx)生成モジュール
pdf_builder.py と同じレイアウト仕様(A4縦/1ページ5組)をWordの表で再現する
さらに「施工中画像・その他」ページ(1ページ最大20枚、4枚×5段)にも対応する
"""
import os
import math
import tempfile

from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageOps

from layout_utils import (
    PAGE_W_MM,
    PAGE_H_MM,
    MARGIN_MM,
    COMMENT_BOX_H_MM,
    COMMENT_GAP_MM,
    COMMENT_FONT_PT,
    GAP_MM,
    EXTRA_PHOTOS_PER_ROW,
    EXTRA_GAP_MM,
    EXTRA_COL_W_MM,
    EXTRA_ROWS_PER_PAGE,
    EXTRA_PHOTOS_PER_PAGE,
)

# ---- レイアウト定数 (mm) : pdf_builder.py と揃えている ----
TITLE_H_MM = 15
LABEL_H_MM = 8
PAIRS_PER_PAGE = 5

# ---- 画像圧縮の設定 ----
IMAGE_DPI = 200
JPEG_QUALITY = 85

# ---- 表紙に記載する会社情報(固定文言) ----
COMPANY_NAME = "株式会社インクコーポレーション"
COMPANY_ADDRESS = "住所：東京都葛飾区立石8-39-6"
COMPANY_TEL_FAX = "TEL：03-3697-9889　FAX：03-3697-9868"
LOGO_PATH = os.path.join(os.path.dirname(__file__), "assets", "logo.png")

MM_TO_PT = 2.83465


def _load_image(img_path):
    """画像を読み込み、EXIFの向き情報をピクセルに反映してから返す
    (スマホ写真が回転して埋め込まれるのを防ぐ)"""
    im = Image.open(img_path)
    im = ImageOps.exif_transpose(im)
    if im.mode not in ("RGB", "L"):
        im = im.convert("RGB")
    return im


def _fit_size(size_or_im, max_w, max_h):
    """画像の縦横比を保ったまま、指定した枠(max_w x max_h)に収まるサイズを計算する"""
    w, h = size_or_im.size if hasattr(size_or_im, "size") else size_or_im
    ratio = min(max_w / w, max_h / h)
    return w * ratio, h * ratio


def _prepare_image_for_docx(img_path, max_w_mm, max_h_mm, tmp_files,
                             dpi=IMAGE_DPI, quality=JPEG_QUALITY):
    """画像を読み込み、向き補正(EXIF)をした上で、実際にWord上へ配置される
    サイズ(mm)を基準に、印刷でも十分きれいに見える解像度(dpi)まで縮小し、
    JPEGとして一時ファイルに保存する(ファイルサイズを大幅に削減するため)。

    戻り値: (一時ファイルパス, (表示幅mm, 表示高さmm))
    """
    im = _load_image(img_path)
    disp_w, disp_h = _fit_size(im, max_w_mm, max_h_mm)

    target_px_w = max(1, round(disp_w / 25.4 * dpi))
    target_px_h = max(1, round(disp_h / 25.4 * dpi))
    if im.size[0] > target_px_w or im.size[1] > target_px_h:
        im = im.resize((target_px_w, target_px_h), Image.LANCZOS)

    tmp = tempfile.NamedTemporaryFile(suffix=".jpg", delete=False)
    im.save(tmp.name, format="JPEG", quality=quality, optimize=True)
    tmp.close()
    tmp_files.append(tmp.name)
    im.close()

    return tmp.name, (disp_w, disp_h)


def _set_vertical_center(cell):
    """表のセル内の文字・画像を上下中央に配置する"""
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    tcPr.append(vAlign)


def _set_row_widths(row, widths_mm):
    for cell, w in zip(row.cells, widths_mm):
        cell.width = Mm(w)


def _add_centered_picture(paragraph, img_path, width_mm, height_mm):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(img_path, width=Mm(width_mm), height=Mm(height_mm))


def _add_paragraph_bottom_border(paragraph):
    """段落の下部に罫線(記入欄の下線)を引く"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B4B4B4")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_comment_area(cell, box_h_mm, comment_text=""):
    """写真の下に記入欄(入力されたコメント+下線)を追加する"""
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(COMMENT_GAP_MM * MM_TO_PT)
    if comment_text:
        run = p.add_run(comment_text)
        run.font.size = Pt(COMMENT_FONT_PT)
    else:
        # 空でも罫線の高さを保つため、幅ゼロの見えない文字を入れておく
        run = p.add_run(" ")
        run.font.size = Pt(COMMENT_FONT_PT)
    _add_paragraph_bottom_border(p)

    label_line_mm = 5
    remaining_mm = box_h_mm - label_line_mm
    if remaining_mm > 0:
        p.paragraph_format.space_after = Pt(remaining_mm * MM_TO_PT)


def _add_cover_page(doc, property_name, cover_photo_path, contact_name, tmp_img_files):
    """表紙ページ(1ページ目)を作成する

    レイアウトの考え方(pdf_builder.pyと揃えている)：
    - 写真をA4用紙の上下中央に配置する
    - 「写真報告書」は写真の3行分上、「物件名」はさらにその2行分上
    - ロゴ・会社情報は写真の3行分下
    """
    usable_w = PAGE_W_MM - 2 * MARGIN_MM
    LINE_MM = 8
    line_pt = LINE_MM * MM_TO_PT

    BASE_MAX_H_MM = 150
    max_w = usable_w * 0.72
    max_h = BASE_MAX_H_MM * 0.72
    if cover_photo_path:
        photo_path, (cw, ch) = _prepare_image_for_docx(cover_photo_path, max_w, max_h, tmp_img_files)
    else:
        photo_path = None
        cw, ch = 0, max_h * 0.7

    if os.path.exists(LOGO_PATH):
        with Image.open(LOGO_PATH) as logo_im:
            lw, lh = _fit_size(logo_im, 110, 20)
    else:
        lw, lh = 0, 0

    title_line_mm = 9
    subtitle_line_mm = 9.5
    company_block_mm = 7 + 5 + 5
    if contact_name.strip():
        company_block_mm += 6

    total_block_mm = (
        title_line_mm + 2 * LINE_MM + subtitle_line_mm + 3 * LINE_MM
        + ch + 3 * LINE_MM + lh + 4 + company_block_mm
    )
    available_mm = PAGE_H_MM - 2 * MARGIN_MM
    center_offset_mm = max(0, (available_mm - total_block_mm) / 2)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(center_offset_mm * MM_TO_PT)
    title_run = title_p.add_run(f"物件名　{property_name}")
    title_run.bold = False
    title_run.underline = True
    title_run.font.size = Pt(20)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_before = Pt(2 * line_pt)
    subtitle_run = subtitle_p.add_run("写真報告書")
    subtitle_run.font.size = Pt(22)

    if cover_photo_path:
        photo_p = doc.add_paragraph()
        photo_p.paragraph_format.space_before = Pt(3 * line_pt)
        photo_p.paragraph_format.space_after = Pt(3 * line_pt)
        _add_centered_picture(photo_p, photo_path, cw, ch)
    else:
        spacer_p = doc.add_paragraph()
        spacer_p.paragraph_format.space_before = Pt(3 * line_pt)
        spacer_p.paragraph_format.space_after = Pt(3 * line_pt + ch * MM_TO_PT)

    if os.path.exists(LOGO_PATH):
        logo_p = doc.add_paragraph()
        logo_p.paragraph_format.space_after = Pt(4)
        _add_centered_picture(logo_p, LOGO_PATH, lw, lh)

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(COMPANY_NAME)
    name_run.font.size = Pt(13)

    if contact_name.strip():
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_p.add_run(f"担当：{contact_name.strip()}")
        contact_run.font.size = Pt(11)

    addr_p = doc.add_paragraph()
    addr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    addr_run = addr_p.add_run(COMPANY_ADDRESS)
    addr_run.font.size = Pt(11)

    tel_p = doc.add_paragraph()
    tel_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tel_run = tel_p.add_run(COMPANY_TEL_FAX)
    tel_run.font.size = Pt(11)


def _add_extra_photos_pages(doc, title, photo_items, tmp_img_files, first_page_break):
    """「施工中画像・その他」ページを作成する(4枚×5段 = 1ページ最大20枚)
    photo_items: [(写真パス, コメント文字列), ...]
    first_page_break: 最初のページの前に改ページを入れるかどうか
    """
    if not photo_items:
        return

    row_h2 = (PAGE_H_MM - 2 * MARGIN_MM - TITLE_H_MM) / EXTRA_ROWS_PER_PAGE
    photo_zone_h2 = row_h2 - COMMENT_GAP_MM - COMMENT_BOX_H_MM
    pad = 2
    widths2 = [EXTRA_COL_W_MM] * EXTRA_PHOTOS_PER_ROW

    total = len(photo_items)
    for page_start in range(0, total, EXTRA_PHOTOS_PER_PAGE):
        title_p = doc.add_paragraph()
        if page_start > 0 or first_page_break:
            title_p.paragraph_format.page_break_before = True
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(16)

        chunk = photo_items[page_start: page_start + EXTRA_PHOTOS_PER_PAGE]
        n_rows = math.ceil(len(chunk) / EXTRA_PHOTOS_PER_ROW)

        table = doc.add_table(rows=0, cols=EXTRA_PHOTOS_PER_ROW)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        for r in range(n_rows):
            row = table.add_row()
            row.height = Mm(row_h2)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            _set_row_widths(row, widths2)
            cells = row.cells

            for c in range(EXTRA_PHOTOS_PER_ROW):
                idx = r * EXTRA_PHOTOS_PER_ROW + c
                if idx >= len(chunk):
                    continue  # 最後の行で埋まらない分は空セルのままにする
                photo_path_src, comment = chunk[idx]

                img_path, (iw, ih) = _prepare_image_for_docx(
                    photo_path_src, EXTRA_COL_W_MM - 2 * pad, photo_zone_h2 - 2 * pad, tmp_img_files
                )
                p_img = cells[c].paragraphs[0]
                _add_centered_picture(p_img, img_path, iw, ih)
                _add_comment_area(cells[c], COMMENT_BOX_H_MM, comment)
                _set_vertical_center(cells[c])


def build_docx(property_name, pairs, output_path, cover_photo_path=None,
               contact_name="", include_cover=True,
               extra_title="施工中画像", extra_photos=None):
    """
    property_name: str  物件名
    pairs: [(施工前画像パス, 施工後画像パス, 施工前コメント, 施工後コメント), ...]
    output_path: 出力する.docxのパス
    cover_photo_path: 表紙に載せる写真のパス(Noneなら写真なし)
    contact_name: 担当者名
    include_cover: 表紙ページを作るかどうか
    extra_title: 「施工中画像・その他」ページの見出し
    extra_photos: [(写真パス, コメント文字列), ...] (Noneまたは空リストならページ自体作らない)
    """
    doc = Document()

    section = doc.sections[0]
    section.page_width = Mm(PAGE_W_MM)
    section.page_height = Mm(PAGE_H_MM)
    section.top_margin = Mm(MARGIN_MM)
    section.bottom_margin = Mm(MARGIN_MM)
    section.left_margin = Mm(MARGIN_MM)
    section.right_margin = Mm(MARGIN_MM)

    usable_w = PAGE_W_MM - 2 * MARGIN_MM
    col_w = (usable_w - GAP_MM) / 2
    row_h = (PAGE_H_MM - 2 * MARGIN_MM - TITLE_H_MM - LABEL_H_MM) / PAIRS_PER_PAGE
    pad = 2
    widths = [col_w, GAP_MM, col_w]

    photo_zone_h = row_h - COMMENT_GAP_MM - COMMENT_BOX_H_MM

    tmp_img_files = []

    try:
        if include_cover:
            _add_cover_page(doc, property_name, cover_photo_path, contact_name, tmp_img_files)

        total = len(pairs)
        for page_start in range(0, total, PAIRS_PER_PAGE):
            title_p = doc.add_paragraph()
            if page_start > 0 or include_cover:
                title_p.paragraph_format.page_break_before = True
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_p.add_run(property_name)
            title_run.bold = True
            title_run.font.size = Pt(18)

            table = doc.add_table(rows=1, cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False

            header_cells = table.rows[0].cells
            _set_row_widths(table.rows[0], widths)

            hp0 = header_cells[0].paragraphs[0]
            hp0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hr0 = hp0.add_run("施工前")
            hr0.bold = True
            hr0.font.size = Pt(12)

            hp2 = header_cells[2].paragraphs[0]
            hp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hr2 = hp2.add_run("施工後")
            hr2.bold = True
            hr2.font.size = Pt(12)

            chunk = pairs[page_start: page_start + PAIRS_PER_PAGE]
            for pair in chunk:
                before, after = pair[0], pair[1]
                before_comment = pair[2] if len(pair) > 2 else ""
                after_comment = pair[3] if len(pair) > 3 else ""

                row = table.add_row()
                row.height = Mm(row_h)
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                _set_row_widths(row, widths)
                cells = row.cells

                before_path, (bw, bh) = _prepare_image_for_docx(
                    before, col_w - 2 * pad, photo_zone_h - 2 * pad, tmp_img_files
                )
                p_before = cells[0].paragraphs[0]
                _add_centered_picture(p_before, before_path, bw, bh)
                _add_comment_area(cells[0], COMMENT_BOX_H_MM, before_comment)
                _set_vertical_center(cells[0])

                p_arrow = cells[1].paragraphs[0]
                p_arrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
                arrow_run = p_arrow.add_run("→")
                arrow_run.bold = True
                arrow_run.font.size = Pt(28)
                _set_vertical_center(cells[1])

                after_path, (aw, ah) = _prepare_image_for_docx(
                    after, col_w - 2 * pad, photo_zone_h - 2 * pad, tmp_img_files
                )
                p_after = cells[2].paragraphs[0]
                _add_centered_picture(p_after, after_path, aw, ah)
                _add_comment_area(cells[2], COMMENT_BOX_H_MM, after_comment)
                _set_vertical_center(cells[2])

        if extra_photos:
            _add_extra_photos_pages(
                doc, extra_title or "施工中画像", extra_photos, tmp_img_files,
                first_page_break=(include_cover or total > 0),
            )

        doc.save(output_path)
    finally:
        for f in tmp_img_files:
            try:
                os.remove(f)
            except OSError:
                pass

    return output_path
